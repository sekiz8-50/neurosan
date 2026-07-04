"""Offline zelftest van de VIF-tekstketen (zonder Meta/OpenAI/echte mail).

Draait: VIF (Word) → intake-extractie → copy → SEO → trends → GEO-LLM →
brand-bewaker → ATS-administrateur (dry-run). Genereert eerst een voorbeeld-VIF
als die nog niet bestaat. Zo controleer je de hele tekst-/Tigris-keten lokaal.

Gebruik:   python selftest_vif.py
Met echte AI-teksten:   ANTHROPIC_API_KEY=... python selftest_vif.py
"""
import json
import os

# Dummy-waarden zodat config.py importeert; deze test raakt Meta/OpenAI/mail NIET.
for k, v in {
    "META_ACCESS_TOKEN": "x", "META_AD_ACCOUNT_ID": "0", "META_PAGE_ID": "0",
    "OPENAI_API_KEY": "x", "RESEND_API_KEY": "x", "APPROVAL_TO": "test@example.com",
    "PUBLIC_BASE_URL": "https://automation.tecqgroep.test", "SIGNING_SECRET": "x",
    "TIGRIS_SHARED_SECRET": "x",
}.items():
    os.environ.setdefault(k, v)

import agents
import vif_parser
from tools import salesforce, trends

VOORBEELD = os.path.join(os.path.dirname(__file__), "data", "voorbeeld_vif.docx")


def maak_voorbeeld_vif(path: str) -> None:
    """Schrijft een representatief ingevuld VIF (twee-koloms tabel) weg."""
    from docx import Document
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc = Document()
    doc.add_heading("Vacature-Intake-Formulier", level=1)
    velden = [
        ("Functietitel", "Onderhoudsmonteur"),
        ("Gewenste functie", "Monteur"),
        ("Sector", "Productie / Industrie"),
        ("Vakgebied", "Mechatronica / Onderhoud"),
        ("Plaats", "Eindhoven"),
        ("Postcode", "5651"),
        ("Provincie", "Noord-Brabant"),
        ("Opleidingsniveau", "MBO niveau 3/4"),
        ("Werkervaring", "2-5 jaar"),
        ("Dienstverband", "Fulltime, 40 uur"),
        ("Rijbewijs", "B"),
        ("Salaris", "2800 - 3600 euro per maand"),
        ("Opdrachtgever", "Een toonaangevende voedingsmiddelenproducent"),
        ("Skills", "storingsanalyse, hydrauliek, pneumatiek, elektrotechniek"),
        ("Arbeidsvoorwaarden", "reiskostenvergoeding, 27 vakantiedagen, opleidingsbudget"),
        ("Bijzonderheden", "Start z.s.m., 2 posities, ploegendienst bespreekbaar"),
    ]
    tabel = doc.add_table(rows=0, cols=2)
    for label, waarde in velden:
        cellen = tabel.add_row().cells
        cellen[0].text, cellen[1].text = label, waarde
    doc.save(path)


def toon(titel: str, data) -> None:
    print(f"\n=== {titel} ===")
    print(json.dumps(data, ensure_ascii=False, indent=2) if not isinstance(data, str) else data)


def main() -> None:
    if not os.path.exists(VOORBEELD):
        maak_voorbeeld_vif(VOORBEELD)
        print(f"[fixture] voorbeeld-VIF aangemaakt: {VOORBEELD}")

    print(f"[AI] ANTHROPIC_API_KEY {'gezet → LLM-paden' if os.environ.get('ANTHROPIC_API_KEY') else 'leeg → deterministische fallbacks'}")

    raw = vif_parser.parse_vif(VOORBEELD)
    toon("1. Ruwe VIF-tekst", raw)

    vac = agents.vif_to_vacancy(raw)
    vac.setdefault("id", "VIF-selftest")
    toon("2. Intake-extractie (vif_to_vacancy)", vac)

    copy = agents.copy_specialist(vac)
    vac["omschrijving"] = copy.get("omschrijving", {})
    vac["quote"] = copy.get("quote", "")
    toon("3. Copy-specialist (omschrijvingsblokken)", copy)

    seo = agents.seo_specialist(vac)
    vac["seo"], vac["keywords"] = seo, seo.get("keywords", [])
    vac["vacature_url"] = f"https://www.maintec.nl/vacatures/{seo.get('slug')}"
    toon("4. SEO-specialist", seo)

    vac["trends"] = trends.popularity(vac.get("titel", ""), vac.get("plaats", ""))
    toon("5. Google-Trends-specialist", vac["trends"])

    geo = agents.geo_llm_specialist(vac, seo)
    vac["schema_org"], vac["faq"] = geo["schema_org"], geo["faq"]
    toon("6. GEO-LLM-specialist (schema.org + FAQ)", geo)

    vac["review_vacature"] = agents.brand_bewaker(vac)
    toon("7. Brand-bewaker", vac["review_vacature"])

    vac["foto_url"] = "https://automation.tecqgroep.test/beeld/VIF-selftest.png"
    sf = salesforce.create_vacancy(vac)
    toon("8. ATS-administrateur → Salesforce-resultaat", sf)

    print("\n✅ VIF-tekstketen volledig doorlopen (Meta/OpenAI/mail bewust overgeslagen in deze zelftest).")


if __name__ == "__main__":
    main()
