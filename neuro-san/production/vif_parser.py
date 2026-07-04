"""VIF-parser — leest het Vacature-Intake-Formulier (Word) uit.

Sales vult de VIF in met de klant en uploadt 'm op /vif. Hier halen we de
ruwe tekst eruit (paragrafen + tabelcellen). De feitelijke 'wat staat erin?'
-interpretatie doet de LLM in agents.vif_to_vacancy(); deze module blijft dom
en robuust, zodat élk VIF-sjabloon werkt zonder vaste veldposities.
"""
import io

from docx import Document


def _extract(doc: "Document") -> str:
    """Alle tekst uit het document als platte, regelgescheiden string."""
    regels: list[str] = []

    for para in doc.paragraphs:
        tekst = para.text.strip()
        if tekst:
            regels.append(tekst)

    for tabel in doc.tables:
        for rij in tabel.rows:
            # ontdubbel samengevoegde cellen (Word herhaalt de tekst per merge)
            uniek: list[str] = []
            for cel in (c.text.strip() for c in rij.cells):
                if cel and (not uniek or uniek[-1] != cel):
                    uniek.append(cel)
            if not uniek:
                continue
            # twee-koloms VIF-tabellen → 'label: waarde'
            regels.append(f"{uniek[0]}: {' | '.join(uniek[1:])}" if len(uniek) >= 2 else uniek[0])

    return "\n".join(regels)


def _pdf_form_fields(data: bytes) -> dict:
    """AcroForm-veldwaarden uit een interactieve (niet-platgeslagen) PDF."""
    try:
        from pypdf import PdfReader
        velden = PdfReader(io.BytesIO(data)).get_fields() or {}
        return {k: str(v.get("/V")).strip() for k, v in velden.items()
                if v.get("/V") not in (None, "", "/Off")}
    except Exception as e:
        print(f"[vif-pdf] AcroForm-velden overgeslagen: {e}")
        return {}


def _extract_pdf(data: bytes) -> str:
    """Tekst + tabellen (tekstlaag) én AcroForm-veldwaarden uit een PDF.
    Geen OCR — een gescande PDF zonder tekstlaag geeft weinig tekst."""
    import pdfplumber
    regels: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            tekst = (page.extract_text() or "").strip()
            if tekst:
                regels.append(tekst)
            for tabel in (page.extract_tables() or []):
                for rij in tabel:
                    cellen = [c.strip() for c in rij if c and c.strip()]
                    if len(cellen) >= 2:
                        regels.append(f"{cellen[0]}: {' | '.join(cellen[1:])}")
    for label, waarde in _pdf_form_fields(data).items():
        regels.append(f"{label}: {waarde}")
    return "\n".join(regels)


def parse_vif(path: str) -> str:
    """Leest de VIF vanaf een pad — Word (.docx) of PDF (.pdf)."""
    if path.lower().endswith(".pdf"):
        with open(path, "rb") as f:
            return _extract_pdf(f.read())
    return _extract(Document(path))


def parse_vif_bytes(data: bytes, is_pdf: bool = False) -> str:
    """Leest de VIF vanuit ruwe bytes (handig voor een HTTP-upload)."""
    return _extract_pdf(data) if is_pdf else _extract(Document(io.BytesIO(data)))
