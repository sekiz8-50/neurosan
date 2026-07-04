"""
Bouwt het pagina-pakket voor /zzp-naar-detachering/ (Maintec)
Output: Maintec-Paginapakket-ZZP-naar-Detachering.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MAINTEC_ORANGE = RGBColor(0xE8, 0x5A, 0x1F)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_BG = RGBColor(0xF5, 0xF5, 0xF5)

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# --- Default style ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_heading(text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_para(text, bold=False, italic=False, color=None, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p

def add_bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

def add_numbered(items):
    for it in items:
        doc.add_paragraph(it, style='List Number')

def add_label(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = MAINTEC_ORANGE
    return p

def add_divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ============================================================
# COVER
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("PAGINA-PAKKET")
r.font.size = Pt(11); r.font.color.rgb = MAINTEC_ORANGE; r.bold = True

t2 = doc.add_paragraph()
r = t2.add_run("/zzp-naar-detachering/")
r.font.size = Pt(28); r.bold = True; r.font.color.rgb = DARK

t3 = doc.add_paragraph()
r = t3.add_run("Opvangnet-pagina voor technische zzp'ers — DBA-handhaving 2026")
r.font.size = Pt(14); r.font.color.rgb = GRAY; r.italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run("Voor: ").bold = True
meta.add_run("Maintec (TecqGroep)\n")
meta.add_run("Bestemd voor: ").bold = True
meta.add_run("Marketingteam, CMS-redactie, webbouwer\n")
meta.add_run("Versie: ").bold = True
meta.add_run("1.0 — mei 2026\n")
meta.add_run("Status: ").bold = True
meta.add_run("Implementatie-klaar (copy + wireframe + tech-specs)")

doc.add_paragraph()
intro = doc.add_paragraph()
intro.add_run(
    "Dit document bevat alles wat nodig is om de pagina /zzp-naar-detachering/ live te zetten: "
    "strategische onderbouwing, wireframe per sectie, volledige Nederlandstalige copy, "
    "SEO-metadata, formulier-specificaties, en een implementatie-checklist. "
    "De pagina is bedoeld als 'opvangnet' voor zzp'ers in de technische sector die door de "
    "DBA-handhaving per 1 januari 2026 hun positie heroverwegen — en als kanaal om hen om te "
    "zetten naar gedetacheerde medewerkers van Maintec."
).italic = True

doc.add_page_break()

# ============================================================
# 1. STRATEGISCHE CONTEXT
# ============================================================
add_heading("1. Strategische context", 1, MAINTEC_ORANGE)

add_heading("1.1 Aanleiding", 2)
doc.add_paragraph(
    "Per 1 januari 2026 begint de Belastingdienst met reguliere handhaving op schijnzelfstandigheid "
    "(Wet DBA). Hoewel er een 'zachte landing' geldt — geen verzuimboetes in 2026 — kunnen "
    "naheffingen, correcties en boetes bij kwaadwillendheid wél worden opgelegd. Opdrachtgevers in "
    "de techniek zijn risicomijdend en stoppen massaal zzp-constructies of zetten ze om naar "
    "detachering en uitzending. De markt verschuift zichtbaar."
)
doc.add_paragraph(
    "Concurrenten YER en DPA hebben al een specifieke landingspagina voor deze doelgroep. "
    "Maintec heeft hier momenteel niets over staan op de website — terwijl Maintec juist het "
    "ideale alternatief biedt voor de technische zzp'er: detachering met behoud van tarief, "
    "afwisseling en vakmanschap, plus de zekerheid van een werkgever."
)

add_heading("1.2 Doelgroep", 2)
add_bullets([
    "Technische zzp'ers in machinebouw, scheepsbouw, automotive, civiel/utiliteit, fietsindustrie",
    "Functies: monteur, lasser, machinist/CNC-operator, elektromonteur, werkvoorbereider, engineer, projectleider",
    "Leeftijd: 25-55, mix van vakkrachten en hoger opgeleide engineers",
    "Pijnpunten: opdracht stopgezet of dreigt te stoppen, onzekerheid over fiscale positie, vrees voor inkomensverlies, weerstand tegen 'klassiek dienstverband'",
    "Waar zoeken ze: Google ('schijnzelfstandigheid techniek', 'alternatief zzp 2026', 'detachering technisch'), LinkedIn, vakcommunities",
])

add_heading("1.3 Pagina-doel & KPI's", 2)
doc.add_paragraph("Primair doel: leads genereren van technische zzp'ers die overwegen over te stappen naar detachering.")
add_label("KPI'S OM TE MONITOREN")
add_bullets([
    "Conversieratio bezoeker → ingevulde intake (target: 4-7%)",
    "Conversieratio bezoeker → tariefcalculator gebruikt (target: 12-18%)",
    "Conversieratio calculator-gebruiker → e-mail achtergelaten (target: 35-45%)",
    "Organische posities op zoektermen: 'zzp naar detachering', 'alternatief zzp 2026', 'schijnzelfstandigheid techniek', 'detachering monteur/lasser/engineer'",
    "Aantal intakes per maand (target: 30-50 in de eerste 6 maanden)",
])

add_heading("1.4 Boodschap-architectuur (positionering)", 2)
add_label("KERNBOODSCHAP")
doc.add_paragraph(
    '"Je hoeft niet terug naar een vast contract. Detachering bij Maintec geeft je vrijheid van afwisseling, '
    'behoud van een sterk uurtarief én de zekerheid van een werkgever — zonder DBA-risico."',
).runs[0].italic = True

add_label("ONDERSTEUNENDE BOODSCHAPPEN")
add_bullets([
    "Behoud je tarief: detachering hoeft niet te betekenen dat je inkomen daalt (toon dit met de calculator)",
    "Behoud je vrijheid: wisselende opdrachten, eigen specialisme, ruimte voor groei",
    "Win zekerheid: doorbetaling bij ziekte, pensioen, scholing, geen administratie",
    "Maintec is geen klassiek uitzendbureau — wij zijn een werkgever met technisch DNA, NEN-4400 en VCU gecertificeerd",
])

add_label("TONE OF VOICE")
add_bullets([
    "Direct en eerlijk — geen verkooppraat, wel duidelijke voordelen",
    "Empathisch — erkennen dat de overstap emotioneel beladen is",
    "Vakgericht — taal van de werkvloer, niet HR-jargon",
    "Tutoyeren ('je', 'jij'), zoals de rest van maintec.nl",
])

doc.add_page_break()

# ============================================================
# 2. TECHNISCHE SPECIFICATIES (SEO + URL)
# ============================================================
add_heading("2. Technische specificaties", 1, MAINTEC_ORANGE)

add_heading("2.1 URL & navigatie", 2)
add_bullets([
    "Primaire URL: https://maintec.nl/zzp-naar-detachering/",
    "Alternatieve URL (indien dieper genest gewenst): /werken-voor-maintec/zzp-naar-detachering/ — beide kunnen via 301 naar de canonical wijzen",
    "Canonical: <link rel=\"canonical\" href=\"https://maintec.nl/zzp-naar-detachering/\" />",
    "Hoofdnavigatie: voeg item toe onder 'Werken voor Maintec' → 'Van zzp naar detachering'",
    "Footer: voeg toe onder 'Werken voor Maintec' — kolom 'Voor kandidaten'",
    "Cross-link vanuit /voor-opdrachtgevers/dba-detachering-2026/ (als die er straks staat) en /nieuws/ DBA-artikelen",
])

add_heading("2.2 SEO-metadata", 2)
add_label("META TITLE (max 60 tekens)")
doc.add_paragraph("Van zzp naar detachering in de techniek | Maintec")

add_label("META DESCRIPTION (max 155 tekens)")
doc.add_paragraph(
    "Stopt je zzp-opdracht door DBA-handhaving? Stap soepel over naar detachering bij Maintec. "
    "Behoud je tarief, win zekerheid. Bereken je netto-inkomen."
)

add_label("PRIMAIRE ZOEKWOORDEN")
add_bullets([
    "zzp naar detachering",
    "alternatief zzp 2026",
    "schijnzelfstandigheid techniek",
    "detachering monteur / lasser / engineer",
    "wet dba 2026 zzp",
])

add_label("OPEN GRAPH / SOCIAL")
add_bullets([
    "og:title: \"Van zzp naar detachering — zonder concessies | Maintec\"",
    "og:description: \"Detachering bij Maintec: behoud van tarief, vrijheid van afwisseling, zekerheid van een werkgever.\"",
    "og:image: nieuw te maken — Maintec-monteur op locatie, oranje accent, tekst-overlay 'DBA-proof werken'",
])

add_heading("2.3 Schema.org structured data", 2)
doc.add_paragraph("Plaats de volgende JSON-LD in de <head>:")
sample = (
    "{\n"
    '  "@context": "https://schema.org",\n'
    '  "@type": "FAQPage",\n'
    '  "mainEntity": [\n'
    "    { /* zie alle FAQ-vragen uit sectie 4.8 */ }\n"
    "  ]\n"
    "}"
)
p = doc.add_paragraph(sample)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(9)

add_heading("2.4 Tracking (GA4 / GTM)", 2)
add_bullets([
    "Event: 'calculator_started' — bij eerste interactie met tariefcalculator",
    "Event: 'calculator_completed' — bij tonen van resultaat",
    "Event: 'lead_form_view' — bij scrollen naar intake-formulier",
    "Event: 'lead_form_submit' — bij succesvolle submit",
    "Event: 'whitepaper_download' — bij download van DBA-speelboek",
    "Event: 'whatsapp_click' / 'phone_click' — bij directe contact-CTAs",
])

doc.add_page_break()

# ============================================================
# 3. WIREFRAME (SECTIE-OPBOUW)
# ============================================================
add_heading("3. Wireframe — sectie-opbouw", 1, MAINTEC_ORANGE)
doc.add_paragraph(
    "De pagina volgt een bewezen conversie-architectuur: empathie → urgentie → propositie → bewijs → "
    "actie. Elke sectie heeft één duidelijk doel. Hieronder per sectie het ASCII-wireframe en de bedoeling."
)

sections_overview = [
    ("01", "Hero", "Erkenning + propositie + dubbele CTA (calculator + intake)"),
    ("02", "Wat verandert er per 2026?", "Korte uitleg DBA-handhaving — autoriteit en context"),
    ("03", "Zzp vs. detachering — eerlijke vergelijking", "Helder tabel, geen verkooppraat"),
    ("04", "Waarom detachering bij Maintec?", "4 redenen, met certificeringen als bewijslast"),
    ("05", "Tariefcalculator (LEAD MAGNET)", "Interactieve tool — bereken je netto-inkomen"),
    ("06", "Hoe werkt de overstap?", "3 stappen, max 2 weken — wegnemen drempel"),
    ("07", "Verhalen van vakmensen", "2-3 mini-cases / quotes (echte mensen, foto)"),
    ("08", "FAQ", "Wegnemen van 8 specifieke bezwaren / vragen"),
    ("09", "Slot-CTA + contact", "Dubbele CTA + WhatsApp + telefoon + recruiter-contactkaart"),
]

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Light Grid Accent 2'
hdr = tbl.rows[0].cells
hdr[0].text = "#"; hdr[1].text = "Sectie"; hdr[2].text = "Doel"
for n, s, p in sections_overview:
    row = tbl.add_row().cells
    row[0].text = n; row[1].text = s; row[2].text = p

doc.add_paragraph()
add_label("VISUEEL WIREFRAME (TOP-NAAR-BENEDEN)")
wire = (
"┌─────────────────────────────────────────────────────────┐\n"
"│  [LOGO]               NAV               [Vacatures] [☰] │\n"
"├─────────────────────────────────────────────────────────┤\n"
"│                                                         │\n"
"│  01 HERO                                                │\n"
"│  ┌─────────────────────────────┬──────────────────────┐ │\n"
"│  │  H1: Stopt je zzp-opdracht? │                      │ │\n"
"│  │  Subtitle (empathie)        │   [FOTO MONTEUR]     │ │\n"
"│  │  [BEREKEN MIJN TARIEF →]    │                      │ │\n"
"│  │  [Plan kennismaking]        │                      │ │\n"
"│  └─────────────────────────────┴──────────────────────┘ │\n"
"│                                                         │\n"
"│  02 WAT VERANDERT ER?  (3 bullets + 1 quote)            │\n"
"│  ───────────────────────────────────────────────────    │\n"
"│  03 ZZP vs DETACHERING — vergelijkingstabel             │\n"
"│  ┌──────────────┬──────────────┬───────────────────────┐│\n"
"│  │              │   Zzp        │   Detachering Maintec ││\n"
"│  │  Tarief      │   €X         │   €X (vergelijkbaar)  ││\n"
"│  │  Ziekte      │   ...        │   ...                 ││\n"
"│  │  Pensioen    │   ...        │   ...                 ││\n"
"│  └──────────────┴──────────────┴───────────────────────┘│\n"
"│  04 WAAROM MAINTEC  (4 cards in 2x2 grid)               │\n"
"│  [CARD 1] [CARD 2]                                      │\n"
"│  [CARD 3] [CARD 4]                                      │\n"
"│                                                         │\n"
"│  05 ★ TARIEFCALCULATOR (LEAD MAGNET) ★                  │\n"
"│  ┌─────────────────────────────────────────────────────┐│\n"
"│  │  Stap 1/3: Wat is je huidige uurtarief?             ││\n"
"│  │  [          €__,__/uur       ]                      ││\n"
"│  │  [Volgende →]                                       ││\n"
"│  └─────────────────────────────────────────────────────┘│\n"
"│  → resultaat: netto-vergelijking + 'mail mij dit resultaat'│\n"
"│                                                         │\n"
"│  06 HOE WERKT DE OVERSTAP? (3 stappen-grafiek)          │\n"
"│  [1 INTAKE]  →  [2 MATCH]  →  [3 START]                 │\n"
"│                                                         │\n"
"│  07 VERHALEN — testimonial-carrousel                    │\n"
"│  ───────────────────────────────────────────────────    │\n"
"│  08 FAQ — accordion (8 vragen)                          │\n"
"│  ───────────────────────────────────────────────────    │\n"
"│  09 SLOT-CTA — split (formulier links, contactblok rechts)│\n"
"│  [INTAKE-FORMULIER]  |  [WhatsApp] [Bellen] [Recruiter] │\n"
"└─────────────────────────────────────────────────────────┘\n"
)
p = doc.add_paragraph(wire)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(8)

doc.add_page_break()

# ============================================================
# 4. VOLLEDIGE COPY (PER SECTIE)
# ============================================================
add_heading("4. Volledige copy", 1, MAINTEC_ORANGE)
doc.add_paragraph(
    "Hieronder de complete copy per sectie. CMS-redactie kan dit één-op-één overnemen. "
    "Vetgedrukte stukken zijn in te zetten als nadruk. Plaatshouders voor afbeeldingen, "
    "bedragen of namen staan tussen vierkante haken."
).italic = True

# ---------- 4.1 HERO ----------
add_heading("4.1 Hero", 2)
add_label("H1")
doc.add_paragraph("Stopt je zzp-opdracht door DBA? Stap over zonder concessies.")
add_label("SUBTITLE")
doc.add_paragraph(
    "Per 1 januari 2026 handhaaft de Belastingdienst weer op schijnzelfstandigheid. "
    "Opdrachtgevers in de techniek stoppen massaal zzp-contracten. Bij Maintec hou je je tarief, "
    "je vakgebied en je vrijheid — en krijg je er de zekerheid van een werkgever bij."
)
add_label("PRIMAIRE CTA (knop)")
doc.add_paragraph("Bereken mijn netto-tarief →")
add_label("SECUNDAIRE CTA (link)")
doc.add_paragraph("Plan een vrijblijvend kennismakingsgesprek")
add_label("VISUEEL")
doc.add_paragraph(
    "Hero-image: technicus op locatie (bij voorkeur uit eigen plaatsingen — Maintec-foto), "
    "natuurlijk licht, oranje accent. Alt-text: 'Technicus aan het werk via Maintec-detachering'."
)

# ---------- 4.2 WAT VERANDERT ER ----------
add_heading("4.2 Wat verandert er per 1 januari 2026?", 2)
add_label("H2")
doc.add_paragraph("De zachte landing is voorbij — wat betekent dat voor jou?")
add_label("BODY")
doc.add_paragraph(
    "Sinds 1 januari 2026 handhaaft de Belastingdienst weer actief op de Wet DBA. "
    "In 2026 worden nog geen verzuimboetes opgelegd ('zachte landing'), maar naheffingen, "
    "correctieverplichtingen en boetes bij kwaadwillendheid blijven mogelijk. Opdrachtgevers willen "
    "geen risico lopen en stoppen daarom met veel zzp-contracten."
)
add_label("3 BULLETS")
add_bullets([
    "Opdrachtgevers in machinebouw, automotive en utiliteit zetten lopende zzp-opdrachten om naar detachering of beëindigen ze.",
    "Engineering- en montagebureaus eisen vaker een werkgeversconstructie achter de vakman.",
    "De wachttijd op een nieuwe zzp-klus loopt op — terwijl gedetacheerde collega's gewoon doorwerken.",
])
add_label("QUOTE (in highlight-box)")
p = doc.add_paragraph(
    '"De vraag naar onze technische detachering is sinds eind 2025 met [XX]% gestegen. '
    'Veel vakmensen die we nu binnenhalen, kwamen eerst als zzp\'er bij dezelfde opdrachtgever."'
)
p.runs[0].italic = True
add_para("— [Naam], recruitmentmanager Maintec", italic=True, color=GRAY, size=10)

# ---------- 4.3 VERGELIJKING ----------
add_heading("4.3 Zzp vs. detachering bij Maintec — eerlijk vergeleken", 2)
add_label("H2")
doc.add_paragraph("Wat je inlevert. Wat je terugkrijgt.")
add_label("INTRO")
doc.add_paragraph(
    "We doen niet alsof detachering hetzelfde is als zzp'en. Het verschilt op een paar punten — "
    "en op andere punten win je juist. Hieronder de eerlijke vergelijking:"
)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Light Grid Accent 2'
hdr = tbl.rows[0].cells
hdr[0].text = ""
hdr[1].text = "Als zzp'er"
hdr[2].text = "Gedetacheerd bij Maintec"
rows_data = [
    ("Uurtarief / inkomen",
     "Je factureert het bruto-tarief; netto verschilt sterk per situatie",
     "Marktconform bruto-uurloon + vakantiegeld + toeslagen (NBBU-cao). Bereken je netto met de calculator hieronder."),
    ("Afwisseling van opdrachten",
     "Je kiest zelf — als je opdrachten kunt vinden",
     "Wij regelen wisselende opdrachten in jouw vakgebied en regio"),
    ("Doorbetaling bij ziekte",
     "Alleen met een arbeidsongeschiktheidsverzekering (eigen kosten)",
     "100% doorbetaling vanaf dag 1, conform cao"),
    ("Pensioen",
     "Zelf regelen — als je eraan toekomt",
     "Automatische opbouw via verplichte pensioenregeling"),
    ("Vakantie",
     "Geen vakantiegeld, geen doorbetaalde vrije dagen",
     "Vakantiegeld + 25 vakantiedagen per jaar (cao)"),
    ("Administratie",
     "BTW-aangifte, IB-aangifte, urenadministratie, facturen",
     "Geen administratie — wij doen alles"),
    ("DBA-risico",
     "Naheffing bij schijnzelfstandigheid mogelijk",
     "Nul risico — je bent in loondienst bij Maintec"),
    ("Opleiding & certificering",
     "Eigen kosten (€500-€2.500 per cursus)",
     "Maintec investeert in jouw certificeringen (VCA, lasdiploma's, etc.)"),
    ("Vrijheid van keuze",
     "Volledige vrijheid wie je accepteert",
     "Je bespreekt elke nieuwe opdracht met je recruiter — geen verplichting"),
]
for label, zzp, det in rows_data:
    row = tbl.add_row().cells
    row[0].text = label
    row[0].paragraphs[0].runs[0].bold = True
    row[1].text = zzp
    row[2].text = det

doc.add_paragraph()
add_label("CTA ONDER TABEL")
doc.add_paragraph("Bereken zelf je netto-inkomen via detachering →")

# ---------- 4.4 WAAROM MAINTEC ----------
add_heading("4.4 Waarom detachering juist bij Maintec?", 2)
add_label("H2")
doc.add_paragraph("Een werkgever die jouw vak begrijpt — niet 'zomaar' een uitzendbureau.")
add_label("INTRO")
doc.add_paragraph(
    "Maintec is sinds 2001 specialist in technische detachering en uitzending. Geen brede HR-keten, "
    "maar 10 lokale vestigingen met recruiters die zelf uit de techniek komen. Vier redenen waarom "
    "vakmensen voor ons kiezen:"
)

cards = [
    ("Technisch DNA",
     "Onze recruiters spreken de taal van de werkvloer. Geen abstracte functieprofielen — "
     "wij begrijpen het verschil tussen TIG en MIG, tussen onderhoud en revisie, tussen "
     "een werkvoorbereider en een uitvoerder."),
    ("Lokaal netwerk",
     "10 vestigingen door heel Nederland (hoofdkantoor in Barendrecht). We plaatsen "
     "vakkrachten lokaal via ons 'Local for Local'-concept — minder reistijd, meer thuiswerk-balans."),
    ("Vakontwikkeling",
     "Via Maintec Vakschool investeren we in opleidingen en certificeringen. "
     "BBL-trajecten voor groeiers, bijscholing voor ervaren vakmensen. Op onze rekening."),
    ("Gecertificeerd zonder gedoe",
     "NEN-4400, VCU en VRO geregistreerd. Alle wettelijke, fiscale en veiligheidszaken "
     "zijn waterdicht geregeld — voor jou én voor de opdrachtgever waar je werkt."),
]
tbl = doc.add_table(rows=2, cols=2)
tbl.style = 'Light Shading Accent 2'
for i, (h, b) in enumerate(cards):
    cell = tbl.cell(i // 2, i % 2)
    cell.paragraphs[0].add_run(h).bold = True
    cell.add_paragraph(b)

# ---------- 4.5 CALCULATOR ----------
add_heading("4.5 ★ Lead magnet: Tariefcalculator zzp → detachering", 2)
add_label("H2")
doc.add_paragraph("Bereken in 60 seconden wat jij netto overhoudt via detachering")
add_label("INTRO")
doc.add_paragraph(
    "De grootste vraag van technische zzp'ers: 'Lever ik niet in?' "
    "Vul je huidige uurtarief in en zie direct hoe je netto-inkomen zich verhoudt tot een "
    "vergelijkbaar Maintec-detacheringscontract — inclusief vakantiegeld, doorbetaling bij ziekte en pensioenopbouw."
)
add_label("CALCULATOR-STAPPEN (UX)")
add_numbered([
    "STAP 1 — Huidig zzp-uurtarief (€ slider, 35-150)",
    "STAP 2 — Gemiddeld declarabele uren per week (slider, 20-50)",
    "STAP 3 — Vakgebied (dropdown: lassen / mechanisch onderhoud / elektrotechniek / engineering / werkvoorbereiding / overig)",
    "TUSSENSTAP — Toon vergelijking: 'Jouw geschatte netto zzp' vs. 'Maintec netto-equivalent (incl. emolumenten)'",
    "STAP 4 (LEAD CAPTURE) — 'Mail mij dit resultaat + persoonlijke berekening' (e-mail + telefoon optioneel)",
    "BEDANKEN — 'Een recruiter belt je binnen 1 werkdag. Of bekijk gelijk passende opdrachten →'",
])
add_label("REKENMETHODE (BACKEND)")
doc.add_paragraph(
    "Hanteer NBBU-cao-uurloon op basis van vakgebied en ervaring. "
    "Tel daarbij op: 8% vakantiegeld + pensioenopbouw (~7% werkgeverslast equivalent) + "
    "25 vakantiedagen (omgerekend ~9,6% van het bruto-jaarloon) + waarde van doorbetaling bij "
    "ziekte (gemiddeld 2,4% jaarinkomen). Trek bij zzp af: gemiddelde AOV-premie, "
    "pensioeneigen-inleg, administratiekosten, schatting belastingdruk. "
    "Toon altijd 'Indicatie — exacte berekening op intake'."
)
add_label("DESIGN-CRITERIA")
add_bullets([
    "Maximaal 3 visuele stappen vóór lead capture — niet te lang",
    "Resultaat altijd tonen ook bij niet-ingevuld e-mailadres (geen 'gating' van data)",
    "Mobiel-first: sliders moeten daar werken",
    "Trust-elementen onder formulier: 'Maintec gebruikt je gegevens alleen voor dit advies — privacy statement'",
])

# ---------- 4.6 HOE WERKT DE OVERSTAP ----------
add_heading("4.6 Hoe werkt de overstap?", 2)
add_label("H2")
doc.add_paragraph("Drie stappen. Maximaal twee weken.")
add_numbered([
    "INTAKE (30 minuten) — telefonisch of op een vestiging bij jou in de buurt. "
    "We bespreken je vakgebied, je huidige tarief en je voorkeuren qua opdrachten en regio.",
    "MATCH — binnen een week zoeken we 1-3 passende opdrachten uit ons netwerk. "
    "Jij kiest. Geen verplichting.",
    "CONTRACT & START — getekend en je staat op de werkvloer. "
    "Eerste salaris uiterlijk binnen 30 dagen. Je oude zzp-opdracht kunnen we vaak omzetten "
    "naar een detacheringscontract bij dezelfde opdrachtgever.",
])
add_label("REASSURANCE")
doc.add_paragraph(
    "Twijfel je nog? Plan een vrijblijvende oriëntatie zonder verplichting. "
    "Veel vakmensen die we plaatsen, kwamen eerst voor 'gewoon even sparren'."
)

# ---------- 4.7 TESTIMONIALS ----------
add_heading("4.7 Verhalen van vakmensen die de overstap maakten", 2)
add_label("H2")
doc.add_paragraph("Zij gingen je voor — en bleven.")
add_label("INSTRUCTIE VOOR CMS")
doc.add_paragraph(
    "Werf 3 testimonials van Maintec-vakmensen die in 2025/2026 vanuit zzp zijn overgestapt. "
    "Vraag toestemming voor foto + voor- en achternaam + functie + werkgebied. "
    "Hieronder template-quotes als plaatsvulling tot echte verhalen binnenkomen:"
)
add_label("TEMPLATE QUOTE 1 — LASSER, 38 JAAR")
p = doc.add_paragraph(
    '"Ik dacht echt dat ik er financieel op achteruit zou gaan. Met de calculator zag ik dat het '
    'maar [X]% scheelde — en daar krijg ik vakantiegeld, pensioen en een opleidingsbudget voor terug. '
    'De rust die dat geeft is veel waard."'
)
p.runs[0].italic = True
add_para("— [Naam], TIG-lasser bij [opdrachtgever] via Maintec", italic=True, color=GRAY, size=10)

add_label("TEMPLATE QUOTE 2 — ENGINEER, 45 JAAR")
p = doc.add_paragraph(
    '"Mijn opdrachtgever stopte met zzp\'ers. Ik wilde geen vast contract — dan ben je 5 jaar '
    'aan dezelfde tekentafel vast. Bij Maintec wissel ik elke 6-12 maanden, behoud mijn netwerk '
    'en doe waar ik goed in ben."'
)
p.runs[0].italic = True
add_para("— [Naam], werkvoorbereider machinebouw via Maintec", italic=True, color=GRAY, size=10)

# ---------- 4.8 FAQ ----------
add_heading("4.8 Veelgestelde vragen", 2)
add_label("H2")
doc.add_paragraph("Eerlijke antwoorden op de vragen die we het meest krijgen.")

faqs = [
    ("Lever ik echt niet in qua tarief?",
     "Vaak minder dan zzp'ers denken. Bij Maintec krijg je marktconform brutoloon (NBBU-cao) "
     "plus vakantiegeld, pensioen, doorbetaling bij ziekte en opleidingsbudget. Tel je dat op, "
     "dan ligt het Maintec-equivalent meestal op 85-100% van het netto zzp-inkomen — en dat is "
     "zonder de uren die je nu kwijt bent aan administratie en acquisitie. Gebruik de calculator "
     "voor jouw situatie."),
    ("Mag ik nog steeds zelf kiezen welke opdracht ik doe?",
     "Ja. We bespreken elke opdracht vóór we je voorstellen. Past het qua vakgebied, reisafstand "
     "of opdrachtgever niet, dan zoeken we verder. Geen verplichting."),
    ("Kan mijn huidige opdrachtgever me overnemen via Maintec?",
     "Vaak wel. Veel opdrachtgevers in de techniek werken al met Maintec. "
     "Bij intake checken we of jouw opdrachtgever bij ons in het netwerk zit — zo niet, dan "
     "kunnen we contact opnemen. Je behoudt je werkplek, alleen je contractvorm verandert."),
    ("Wat als ik tussen opdrachten zit — krijg ik dan nog salaris?",
     "Ja. Bij Maintec heb je een arbeidsovereenkomst. Tussen opdrachten ontvang je gewoon je salaris "
     "(behoudens uitzonderingen die we tijdens intake helder maken). Geen onzekerheid meer over leeg gat."),
    ("Hoe zit het met mijn BV / mijn huidige onderneming?",
     "Die hoef je niet direct op te heffen. Veel vakmensen houden hun BV slapend aan voor eventuele "
     "toekomstige nevenwerkzaamheden. Onze recruiter denkt mee in jouw fiscale traject — wij werken "
     "samen met een aantal vaste boekhouders die hierin gespecialiseerd zijn."),
    ("Verlies ik mijn vrijheid?",
     "Nee — je verandert van type vrijheid. De vrijheid van 'eigen baas zijn' wisselt voor de "
     "vrijheid van afwisseling zonder administratieve last. Veel vakmensen ervaren dat als rust, "
     "niet als verlies."),
    ("Wat doet Maintec anders dan een uitzendbureau?",
     "Twee dingen: technisch DNA (recruiters uit de sector, geen generalisten) en het werkgevers-"
     "perspectief. Bij Maintec ben je niet 'tijdelijk personeel', maar een vast lid van een collectief — "
     "met opleidingstraject, ontwikkelingspad en betrokken recruiter."),
    ("Hoe snel kan ik overstappen?",
     "Gemiddeld 5-10 werkdagen tussen eerste contact en eerste werkdag, als er een passende "
     "opdracht ligt. We hebben momenteel meer dan 250 openstaande opdrachten in de techniek."),
]
for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run("V: " + q)
    r.bold = True
    doc.add_paragraph("A: " + a)
    doc.add_paragraph()

# ---------- 4.9 SLOT-CTA ----------
add_heading("4.9 Slot-CTA + contactblok", 2)
add_label("H2")
doc.add_paragraph("Klaar voor de overstap? Of eerst sparren?")
add_label("INTRO")
doc.add_paragraph(
    "Vul het formulier in voor een vrijblijvende intake binnen 1 werkdag. Of bel/WhatsApp direct met "
    "een recruiter uit jouw regio."
)
add_label("FORMULIER (LINKS)")
add_bullets([
    "Naam (verplicht)",
    "E-mail (verplicht)",
    "Telefoonnummer (optioneel, versnelt contact)",
    "Vakgebied (dropdown: lassen / mechanisch / elektro / engineering / werkvoorbereiding / overig)",
    "Vestiging in de buurt (dropdown — 10 vestigingen + 'maakt niet uit')",
    "Korte toelichting (optioneel textarea, max 500 tekens)",
    "Toestemming privacy statement (verplicht checkbox)",
])
add_label("FORMULIER-KNOP")
doc.add_paragraph("Vraag intake aan →")

add_label("CONTACTBLOK (RECHTS)")
add_bullets([
    "WhatsApp-knop: 'WhatsApp ons direct →' (link naar Maintec WhatsApp-nummer)",
    "Bellen: '085-XXX XXXX' (klik = bel-actie op mobiel)",
    "Mail: 'detachering@maintec.nl'",
    "Foto + naam van een 'recruiter die zelf uit de techniek komt' (vertrouwensanker)",
    "Onder contactblok: '10 vestigingen in Nederland — zoek de dichtstbijzijnde →' (link naar /vestigingen)",
])

doc.add_page_break()

# ============================================================
# 5. LEAD-CAPTURE & FOLLOW-UP
# ============================================================
add_heading("5. Lead-capture & follow-up", 1, MAINTEC_ORANGE)

add_heading("5.1 Lead-bronnen op deze pagina", 2)
add_bullets([
    "Tariefcalculator — primaire conversie (e-mail in stap 4)",
    "Intake-formulier — secundaire conversie (slot-sectie)",
    "Whitepaper-download — tertiaire conversie (optioneel)",
    "WhatsApp / telefoon — directe conversie (telefoongesprekken tellen voor sales-team)",
])

add_heading("5.2 Velden, validatie en datakwaliteit", 2)
add_bullets([
    "Naam: vrij invulveld, min 2 tekens, geen verplichte achternaam (verlaagt drempel)",
    "E-mail: verplicht, valideer formaat client-side + server-side, weiger gratis throwaway-domeinen (configurable)",
    "Telefoon: vrij optioneel, NL-format hint (06-, 010-, etc.)",
    "Vakgebied: dropdown voor segmentatie naar juiste recruiter-team",
    "Vestiging: dropdown — koppel automatisch aan accountmanager van die vestiging",
    "Toelichting: textarea, max 500 tekens, geen verplichting",
])

add_heading("5.3 Routing naar CRM / ATS", 2)
add_bullets([
    "Submit triggert webhook naar Maintec ATS (Carerix / Connexys / OTYS — afhankelijk van huidige stack)",
    "Tag lead met source = 'zzp-naar-detachering' + UTM-parameters voor attributie",
    "Wijs toe aan recruiter op basis van 'vestiging' en 'vakgebied'",
    "SLA: eerste contact binnen 1 werkdag — automatische e-mailbevestiging direct na submit",
    "Bij geen contact binnen 24 uur: escalatie naar teamlead",
])

add_heading("5.4 E-mail follow-up sequence", 2)
add_label("E-MAIL 1 — DIRECT NA SUBMIT (BEVESTIGING)")
add_bullets([
    "Onderwerp: 'Bedankt — een recruiter belt je binnen 1 werkdag'",
    "Body: persoonlijke begroeting, wat ze kunnen verwachten, link naar vestigings-info, recruiter-foto",
    "CTA: 'Bekijk nu al passende opdrachten →'",
])
add_label("E-MAIL 2 — DAG 3 (BIJ GEEN INTAKE-AFSPRAAK)")
add_bullets([
    "Onderwerp: 'Nog niet kunnen praten? Hier is alvast extra info'",
    "Body: case-study + link naar whitepaper DBA-speelboek",
    "CTA: 'Plan zelf een gesprek →' (Calendly-link recruiter)",
])
add_label("E-MAIL 3 — DAG 10 (BIJ GEEN INTAKE)")
add_bullets([
    "Onderwerp: 'Markt verandert — wil je nog steeds zzp blijven?'",
    "Body: korte trend-update + uitnodiging om vrijblijvend terug te bellen",
    "CTA: 'WhatsApp ons →'",
])

doc.add_page_break()

# ============================================================
# 6. IMPLEMENTATIE-CHECKLIST
# ============================================================
add_heading("6. Implementatie-checklist", 1, MAINTEC_ORANGE)

add_heading("6.1 Voor lancering (week 1-2)", 2)
add_bullets([
    "☐ Pagina-skelet in CMS bouwen volgens wireframe (sectie 3)",
    "☐ Copy uit sectie 4 plaatsen — laat marketingmanager checken op tone of voice",
    "☐ Hero-foto laten maken/selecteren (eigen Maintec-fotografie) + alt-tekst",
    "☐ Vergelijkingstabel (sectie 4.3) — controle door HR/payroll op cao-juistheid",
    "☐ Calculator-rekenmodel valideren met finance/payroll (sectie 4.5 — rekenmethode)",
    "☐ Calculator front-end laten bouwen (web-development — schatting 1-2 sprints)",
    "☐ Testimonials werven (3 vakmensen, foto + quote)",
    "☐ FAQ Schema.org JSON-LD plaatsen (sectie 2.3)",
    "☐ Formulier koppelen aan ATS via webhook",
    "☐ E-mail-flows opzetten in marketing automation tool (sectie 5.4)",
    "☐ GA4 / GTM events configureren (sectie 2.4)",
    "☐ Hoofdnavigatie en footer aanpassen",
    "☐ Cross-links plaatsen vanuit /voor-opdrachtgevers/, /nieuws/, /werken-voor-maintec/",
])

add_heading("6.2 Lancering & promotie", 2)
add_bullets([
    "☐ XML-sitemap updaten en bij Google Search Console indienen",
    "☐ LinkedIn-post + 3 follow-up posts in 4 weken (targeting: zzp + techniek)",
    "☐ Google Ads-campagne op kernzoekwoorden (€[budget]/maand startbudget)",
    "☐ Nieuwsbrief naar bestaande database (zzp-tag) — 1 mailing met de calculator als focus",
    "☐ Persbericht naar vakmedia (Metaal & Techniek, Vakblad Energietechniek, etc.)",
    "☐ Briefing aan alle recruiters: hoe leads opvolgen, wat te zeggen, calculator demo",
])

add_heading("6.3 Na lancering — A/B-tests & optimalisatie", 2)
add_bullets([
    "☐ A/B-test hero H1: 'Stopt je zzp-opdracht?' vs. 'Wat is je netto na DBA?'",
    "☐ A/B-test CTA-knop: 'Bereken mijn netto-tarief' vs. 'Wat hou ik over?'",
    "☐ A/B-test positie calculator: hoog (na hero) vs. midden (huidige positie)",
    "☐ Heatmap-analyse na 4 weken (Hotjar / Microsoft Clarity)",
    "☐ Maandelijkse review conversieratio's vs. KPI-targets (sectie 1.3)",
    "☐ Conversie-interview met 5 vakmensen die converteerden — wat gaf de doorslag?",
])

doc.add_page_break()

# ============================================================
# 7. BIJLAGE — VOORBEELD BEKEKEN BIJ CONCURRENTIE
# ============================================================
add_heading("7. Bijlage — concurrentie-referenties", 1, MAINTEC_ORANGE)
doc.add_paragraph(
    "Onderstaande concurrenten hebben al een vergelijkbare pagina. Gebruik ter inspiratie, "
    "maar laat de toon en propositie van Maintec eigen blijven (vakgericht, lokaal, geen HR-jargon)."
)
add_bullets([
    "YER — https://www.yer.nl/kandidaten/zzp-schijnzelfstandigheid/  →  Bevat 10-vragen DBA-check + informatiegids als lead magnet. Sterk in autoriteit (Belastingdienst-toon), zwak in empathie (geen verhalen).",
    "Brunel, DPA — Hebben kortere pagina's, gericht op opdrachtgevers. Minder relevant als directe benchmark voor zzp-pagina.",
    "Bring at Work, Flexhub — Niche-bureaus die de detachering-route stevig in de markt zetten. Goed voor inspiratie qua copy-toon.",
])

add_heading("Slotwoord", 2)
doc.add_paragraph(
    "Deze pagina is een opvangnet voor een doelgroep die nú actief op zoek is. "
    "Goede uitvoering kan binnen 6 maanden tientallen extra plaatsingen opleveren — "
    "een groep die anders bij YER, Brunel of een concurrent landt. "
    "Bij vragen of wijzigingen: vraag een herziene versie aan via je marketingteam."
)

# --- SAVE ---
output_path = "/Users/maintec/CLAUDE26/pagina-pakketten/Maintec-Paginapakket-ZZP-naar-Detachering.docx"
doc.save(output_path)
print(f"OK: {output_path}")
